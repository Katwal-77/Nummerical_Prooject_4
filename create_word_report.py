import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create a new Document
doc = docx.Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_heading('Discovering Laws Behind Real-World Data Using the Least Squares Method', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add author information
author_info = doc.add_paragraph()
author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_info.add_run('Name: Prem Bahadur Katuwal\n').bold = True
author_info.add_run('Student ID: 202424080129\n').bold = True
author_info.add_run('Subject: Numerical Analysis (PhD Level)\n').bold = True
author_info.add_run('Assignment: Use Least Squares Method to Discover the Law Behind Real-World Data\n').bold = True
author_info.add_run('Date: April 19, 2025').bold = True

# Add a page break
doc.add_page_break()

# Add Table of Contents heading
toc_heading = doc.add_heading('Table of Contents', level=1)
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add Table of Contents entries
toc_items = [
    ('1. Introduction', 3),
    ('2. Theoretical Foundation', 3),
    ('   2.1 Mathematical Formulation', 3),
    ('   2.2 Matrix Approach', 4),
    ('   2.3 Normal Equations', 4),
    ('3. Linear Least Squares', 5),
    ('   3.1 Dataset Description', 5),
    ('   3.2 Visualizing the Data', 5),
    ('   3.3 Applying the Method', 6),
    ('   3.4 Results and Interpretation', 7),
    ('4. Polynomial Least Squares', 8),
    ('   4.1 Mathematical Extension', 8),
    ('   4.2 Application to Temperature Data', 8),
    ('   4.3 Model Selection', 9),
    ('5. Non-Linear Least Squares', 10),
    ('   5.1 Exponential Growth Model', 10),
    ('   5.2 Linearization vs. Direct Fitting', 11),
    ('6. Advanced Analysis', 12),
    ('   6.1 Residual Analysis', 12),
    ('   6.2 Confidence Intervals', 13),
    ('   6.3 Goodness of Fit Metrics', 13),
    ('7. Comparison with Other Methods', 14),
    ('8. Conclusion', 15),
    ('9. References', 16)
]

# Create TOC
for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f'{item}').bold = True
    
    # Add tab
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT, docx.enum.text.WD_TAB_LEADER.DOTS)
    
    # Add page number
    p.add_run(f'\t{page}')

# Add a page break
doc.add_page_break()

# Add Introduction
doc.add_heading('1. Introduction', level=1)
intro_text = """
The least squares method is a cornerstone technique in numerical analysis and statistics for discovering underlying patterns and relationships in real-world data. Developed by Carl Friedrich Gauss and Adrien-Marie Legendre in the early 19th century, this method has become indispensable across various scientific disciplines, from physics and engineering to economics and social sciences.

At its core, the least squares method seeks to find the best-fitting curve or mathematical model for a given set of data points by minimizing the sum of the squares of the residuals (the differences between observed values and the values predicted by the model). This approach provides a powerful framework for:

1. Discovering mathematical laws that govern physical phenomena
2. Making predictions based on observed data
3. Understanding relationships between variables
4. Testing hypotheses about underlying mechanisms

This report demonstrates the application of the least squares method to several real-world datasets, detailing the mathematical foundations, implementation procedures, and interpretation of results. Through these examples, we will uncover the hidden laws that govern diverse phenomena and showcase the versatility and power of this numerical technique.
"""
doc.add_paragraph(intro_text)

# Add Theoretical Foundation
doc.add_heading('2. Theoretical Foundation', level=1)

# Mathematical Formulation
doc.add_heading('2.1 Mathematical Formulation', level=2)
math_form_text = """
The fundamental principle of the least squares method is to minimize the sum of squared differences between observed values and the values predicted by a model. Given a set of data points (xi, yi) for i = 1, 2, ..., n, and a model function f(x; β) with parameters β, we aim to find the parameter values that minimize:

S(β) = Σ[yi - f(xi; β)]²

This sum S(β) is called the "sum of squared residuals" or "residual sum of squares" (RSS).
"""
doc.add_paragraph(math_form_text)

# Matrix Approach
doc.add_heading('2.2 Matrix Approach', level=2)
matrix_text = """
For linear models, the least squares problem can be elegantly formulated using matrix notation. If we have a linear model:

f(x; β) = β₁φ₁(x) + β₂φ₂(x) + ... + βₘφₘ(x)

where φⱼ(x) are known functions (e.g., φ₁(x) = x, φ₂(x) = 1 for a straight line), we can define matrices X, y, and β to represent the system.

The sum of squared residuals becomes:

S(β) = ||y - Xβ||² = (y - Xβ)ᵀ(y - Xβ)
"""
doc.add_paragraph(matrix_text)

# Normal Equations
doc.add_heading('2.3 Normal Equations', level=2)
normal_eq_text = """
To find the parameter values β that minimize S(β), we differentiate with respect to β and set the result equal to zero:

∂S(β)/∂β = -2Xᵀ(y - Xβ) = 0

This leads to the normal equations:

XᵀXβ = Xᵀy

The solution is:

β = (XᵀX)⁻¹Xᵀy

assuming XᵀX is invertible. In practice, numerical methods like QR decomposition or Singular Value Decomposition (SVD) are often used to solve this system more stably.
"""
doc.add_paragraph(normal_eq_text)

# Add Linear Least Squares
doc.add_heading('3. Linear Least Squares', level=1)

# Dataset Description
doc.add_heading('3.1 Dataset Description', level=2)
dataset_text = """
We begin with a classic example: the relationship between house size and price. This dataset represents a common application of regression analysis in real estate economics. The data points are as follows:
"""
doc.add_paragraph(dataset_text)

# Add table for house data
table = doc.add_table(rows=11, cols=2)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'House Size (sqft)'
header_cells[1].text = 'Price ($1000s)'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
data = [
    (650, 70),
    (785, 85),
    (1200, 120),
    (1500, 145),
    (1850, 180),
    (2100, 210),
    (2300, 230),
    (2700, 265),
    (3000, 300),
    (3500, 355)
]

for i, (size, price) in enumerate(data, start=1):
    row = table.rows[i].cells
    row[0].text = str(size)
    row[1].text = str(price)

# Visualizing the Data
doc.add_heading('3.2 Visualizing the Data', level=2)
vis_text = """
Below is a scatter plot of the data, illustrating the positive correlation between house size and price:
"""
doc.add_paragraph(vis_text)

# Add scatter plot image
doc.add_picture('scatter_plot.png', width=Inches(6))
scatter_caption = doc.add_paragraph('Figure 1: Scatter plot of house size vs. price')
scatter_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Applying the Method
doc.add_heading('3.3 Applying the Method', level=2)
method_text = """
For a linear model, we seek a relationship of the form:

Price = m × Size + c

where m is the slope and c is the intercept.

Using the matrix formulation described earlier, we can solve the normal equations to find the best-fit parameters. The best-fit parameters obtained are:
- Slope (m): 0.0991
- Intercept (c): 1.9382

Thus, the discovered law is:

Price = 0.0991 × Size + 1.9382 ($1000s)

Or more simply:

Price ≈ 0.10 × Size + 1.94 ($1000s)

The fitted line is shown below:
"""
doc.add_paragraph(method_text)

# Add fitted line image
doc.add_picture('fitted_line.png', width=Inches(6))
fitted_caption = doc.add_paragraph('Figure 2: Linear least squares fit for house size vs. price')
fitted_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Results and Interpretation
doc.add_heading('3.4 Results and Interpretation', level=2)
results_text = """
- The positive slope (0.0991) indicates that each additional square foot of house size corresponds to an increase of approximately $99.1 in the house price.
- The intercept (1.9382) represents the theoretical price when the house size is zero. While not meaningful in this context (as houses cannot have zero size), it serves as the y-intercept for our linear model.
- The coefficient of determination (R²) is 0.9980, indicating that 99.8% of the variance in house prices is explained by the house size alone.
- The fitted line closely follows the trend of the data, suggesting that a linear relationship effectively captures the relationship between house size and price in this dataset.
- This model can be used for price estimation: a 2000 sqft house would be estimated at approximately $199,100 + $1,938 = $201,038.
"""
doc.add_paragraph(results_text)

# Add Polynomial Least Squares
doc.add_heading('4. Polynomial Least Squares', level=1)

# Mathematical Extension
doc.add_heading('4.1 Mathematical Extension', level=2)
poly_math_text = """
While linear models are often sufficient, many real-world phenomena exhibit non-linear relationships. Polynomial regression extends the linear least squares method by using polynomial functions as the basis.

For a polynomial of degree p, the model becomes:

f(x; β) = β₀ + β₁x + β₂x² + ... + βₚxᵖ

The matrix X is modified to include powers of x, but the solution procedure remains the same, using the normal equations.
"""
doc.add_paragraph(poly_math_text)

# Application to Temperature Data
doc.add_heading('4.2 Application to Temperature Data', level=2)
temp_text = """
Let's consider a dataset of average monthly temperatures in a northern hemisphere location:
"""
doc.add_paragraph(temp_text)

# Add table for temperature data
table = doc.add_table(rows=13, cols=2)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Month'
header_cells[1].text = 'Temperature (°C)'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
temp_data = [
    (1, -5.2),
    (2, -3.8),
    (3, 2.4),
    (4, 8.7),
    (5, 14.5),
    (6, 19.2),
    (7, 21.8),
    (8, 20.9),
    (9, 16.3),
    (10, 10.1),
    (11, 4.2),
    (12, -2.5)
]

for i, (month, temp) in enumerate(temp_data, start=1):
    row = table.rows[i].cells
    row[0].text = str(month)
    row[1].text = str(temp)

temp_vis_text = """
This temperature pattern clearly follows a cyclical pattern that cannot be adequately captured by a linear model. The data visualization confirms this:
"""
doc.add_paragraph(temp_vis_text)

# Add temperature data image
doc.add_picture('temperature_data.png', width=Inches(6))
temp_data_caption = doc.add_paragraph('Figure 3: Monthly temperature data')
temp_data_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

poly_fit_text = """
We'll fit polynomials of different degrees to find the best representation. The comparison of different polynomial degrees shows how the fit improves with higher-degree polynomials:
"""
doc.add_paragraph(poly_fit_text)

# Add polynomial comparison image
doc.add_picture('temperature_poly_comparison.png', width=Inches(6))
poly_comp_caption = doc.add_paragraph('Figure 4: Comparison of polynomial fits of different degrees')
poly_comp_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

poly_eq_text = """
For a 4th-degree polynomial, the discovered law is:

Temperature = 0.021x⁴ - 0.595x³ + 4.900x² - 9.137x - 0.417

The fitted curve is shown below:
"""
doc.add_paragraph(poly_eq_text)

# Add temperature fit image
doc.add_picture('temperature_fit.png', width=Inches(6))
temp_fit_caption = doc.add_paragraph('Figure 5: 4th-degree polynomial fit for monthly temperatures')
temp_fit_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Model Selection
doc.add_heading('4.3 Model Selection', level=2)
model_sel_text = """
To determine the optimal polynomial degree, we can use metrics like the coefficient of determination (R²) and adjusted R², which account for model complexity:
"""
doc.add_paragraph(model_sel_text)

# Add R² image
doc.add_picture('temperature_r2.png', width=Inches(6))
r2_caption = doc.add_paragraph('Figure 6: R² and adjusted R² for different polynomial degrees')
r2_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add table for polynomial comparison
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Polynomial Degree'
header_cells[1].text = 'R²'
header_cells[2].text = 'Adjusted R²'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
poly_data = [
    ('1 (linear)', '0.0766', '-0.1286'),
    ('2 (quadratic)', '0.9267', '0.8992'),
    ('3 (cubic)', '0.9601', '0.9372'),
    ('4 (quartic)', '0.9984', '0.9972'),
    ('5 (quintic)', '0.9985', '0.9967'),
    ('6 (sextic)', '0.9985', '0.9959')
]

for i, (degree, r2, adj_r2) in enumerate(poly_data, start=1):
    row = table.rows[i].cells
    row[0].text = degree
    row[1].text = r2
    row[2].text = adj_r2

model_sel_conclusion = """
The 4th-degree polynomial provides an excellent fit with an R² of 0.9984, and higher degrees offer diminishing returns. This suggests that a 4th-degree polynomial effectively captures the seasonal temperature variation.
"""
doc.add_paragraph(model_sel_conclusion)

# Add Non-Linear Least Squares
doc.add_heading('5. Non-Linear Least Squares', level=1)

# Exponential Growth Model
doc.add_heading('5.1 Exponential Growth Model', level=2)
exp_text = """
Some phenomena follow inherently non-linear patterns that cannot be linearized through polynomial transformations. Consider population growth data:
"""
doc.add_paragraph(exp_text)

# Add table for population data
table = doc.add_table(rows=14, cols=2)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Year'
header_cells[1].text = 'Population (billions)'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
pop_data = [
    (1900, 1.65),
    (1910, 1.92),
    (1920, 2.22),
    (1930, 2.59),
    (1940, 2.85),
    (1950, 3.17),
    (1960, 3.65),
    (1970, 4.07),
    (1980, 4.45),
    (1990, 5.31),
    (2000, 6.12),
    (2010, 6.93),
    (2020, 7.79)
]

for i, (year, pop) in enumerate(pop_data, start=1):
    row = table.rows[i].cells
    row[0].text = str(year)
    row[1].text = str(pop)

pop_vis_text = """
The data visualization shows an accelerating growth pattern:
"""
doc.add_paragraph(pop_vis_text)

# Add population data image
doc.add_picture('population_data.png', width=Inches(6))
pop_data_caption = doc.add_paragraph('Figure 7: World population growth data')
pop_data_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

exp_model_text = """
An exponential model is appropriate for population growth:

P(t) = P₀e^(rt)

where P₀ is the initial population and r is the growth rate.
"""
doc.add_paragraph(exp_model_text)

# Linearization vs. Direct Fitting
doc.add_heading('5.2 Linearization vs. Direct Fitting', level=2)
lin_vs_direct_text = """
There are two main approaches to fitting exponential models:

1. Linearization: Taking logarithms of both sides transforms the model into a linear form:
   ln(P(t)) = ln(P₀) + rt
   
   This can then be solved using linear least squares.

2. Direct Non-Linear Fitting: Using iterative methods to directly fit the exponential model to the data.

Both approaches yield similar results for this dataset:
"""
doc.add_paragraph(lin_vs_direct_text)

# Add population growth image
doc.add_picture('population_growth.png', width=Inches(6))
pop_growth_caption = doc.add_paragraph('Figure 8: Exponential growth models for population data')
pop_growth_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

exp_results_text = """
The discovered laws are:
- Linearized Method: P(t) = 1.6966 × e^(0.0127t)
- Direct Non-Linear Method: P(t) = 1.6848 × e^(0.0128t)

where t is years since 1900.

These models indicate a growth rate of approximately 1.27-1.28% per year, with a doubling time of about 54 years. Using these models, we can project future population values:
"""
doc.add_paragraph(exp_results_text)

# Add table for population projections
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Year'
header_cells[1].text = 'Linearized Model (billions)'
header_cells[2].text = 'Non-Linear Model (billions)'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
proj_data = [
    (2030, 8.82, 8.86),
    (2040, 10.02, 10.07),
    (2050, 11.37, 11.44),
    (2060, 12.91, 13.00),
    (2070, 14.65, 14.77)
]

for i, (year, lin, nonlin) in enumerate(proj_data, start=1):
    row = table.rows[i].cells
    row[0].text = str(year)
    row[1].text = str(lin)
    row[2].text = str(nonlin)

# Add Advanced Analysis
doc.add_heading('6. Advanced Analysis', level=1)

# Residual Analysis
doc.add_heading('6.1 Residual Analysis', level=2)
residual_text = """
Residuals are the differences between observed values and values predicted by the model. Analyzing residuals helps assess the goodness of fit and validate model assumptions.

For the house price model, the residuals are:
"""
doc.add_paragraph(residual_text)

# Add table for residuals
table = doc.add_table(rows=11, cols=4)
table.style = 'Table Grid'

# Add header row
header_cells = table.rows[0].cells
header_cells[0].text = 'House Size (sqft)'
header_cells[1].text = 'Actual Price'
header_cells[2].text = 'Predicted Price'
header_cells[3].text = 'Residual'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Add data rows
residual_data = [
    (650, 70, 66.34, 3.66),
    (785, 85, 79.72, 5.28),
    (1200, 120, 120.84, -0.84),
    (1500, 145, 150.57, -5.57),
    (1850, 180, 185.25, -5.25),
    (2100, 210, 210.02, -0.02),
    (2300, 230, 229.84, 0.16),
    (2700, 265, 269.47, -4.47),
    (3000, 300, 299.20, 0.80),
    (3500, 355, 348.74, 6.26)
]

for i, (size, actual, predicted, residual) in enumerate(residual_data, start=1):
    row = table.rows[i].cells
    row[0].text = str(size)
    row[1].text = str(actual)
    row[2].text = str(predicted)
    row[3].text = str(residual)

residual_plot_text = """
The plot below shows the residuals for each data point:
"""
doc.add_paragraph(residual_plot_text)

# Add residuals image
doc.add_picture('residuals.png', width=Inches(6))
residuals_caption = doc.add_paragraph('Figure 9: Residuals for the house price model')
residuals_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

residual_analysis_text = """
Key observations from the residual analysis:
- The residuals are relatively small compared to the actual values, indicating a good fit.
- The residuals appear randomly scattered around zero, with no obvious pattern, suggesting the linear model is appropriate.
- There are no extreme outliers that might unduly influence the fit.

A more comprehensive residual analysis would include:
1. Normality test: Checking if residuals follow a normal distribution
2. Homoscedasticity: Verifying constant variance across the range of predictors
3. Autocorrelation: Testing for independence of residuals
"""
doc.add_paragraph(residual_analysis_text)

# Confidence Intervals
doc.add_heading('6.2 Confidence Intervals', level=2)
conf_int_text = """
Confidence intervals quantify the uncertainty in the estimated parameters and predictions. For the house price model, the 95% confidence intervals for predictions are shown below:
"""
doc.add_paragraph(conf_int_text)

# Add confidence intervals image
doc.add_picture('confidence_intervals.png', width=Inches(6))
conf_int_caption = doc.add_paragraph('Figure 10: Least squares fit with 95% confidence intervals')
conf_int_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

conf_int_desc = """
The gray shaded area represents the 95% confidence interval for the predicted values. Note how the interval widens as we move away from the center of the data, reflecting increased uncertainty in predictions at the extremes.
"""
doc.add_paragraph(conf_int_desc)

# Goodness of Fit Metrics
doc.add_heading('6.3 Goodness of Fit Metrics', level=2)
gof_text = """
Several metrics help evaluate how well the model fits the data:

1. Coefficient of Determination (R²): Measures the proportion of variance explained by the model
   - For the house price model: R² = 0.9980
   - This indicates that 99.80% of the variance in house prices is explained by house size

2. Root Mean Square Error (RMSE): Measures the average magnitude of residuals
   - For the house price model: RMSE = 4.0041
   - This means predictions are off by about $4,004 on average

3. Mean Absolute Error (MAE): Average of absolute residuals
   - For the house price model: MAE = 3.2308
   - This means predictions deviate by about $3,231 on average
"""
doc.add_paragraph(gof_text)

# Add Comparison with Other Methods
doc.add_heading('7. Comparison with Other Methods', level=1)
comparison_text = """
While least squares is powerful, other methods may be more appropriate in certain situations:

1. Robust Regression: Less sensitive to outliers
   - Uses methods like Huber loss or Tukey's bisquare
   - Particularly useful when data contains outliers

2. Ridge Regression: Adds a penalty term to handle multicollinearity
   - Minimizes Σ[yi - f(xi; β)]² + λΣβj²
   - Useful when predictors are highly correlated

3. LASSO Regression: Encourages sparse solutions
   - Minimizes Σ[yi - f(xi; β)]² + λΣ|βj|
   - Useful for feature selection

4. Orthogonal Distance Regression: Accounts for errors in both dependent and independent variables
   - Minimizes perpendicular distances from points to the fitted curve
   - Appropriate when both variables have measurement errors

For our house price example, standard least squares is appropriate because:
- The relationship appears strongly linear
- There are no obvious outliers
- We have a single predictor variable
- The residuals are well-behaved
"""
doc.add_paragraph(comparison_text)

# Add Conclusion
doc.add_heading('8. Conclusion', level=1)
conclusion_text = """
The least squares method provides a powerful framework for discovering mathematical laws that govern real-world phenomena. Through this report, we have:

1. Explored the theoretical foundations of the method, from its mathematical formulation to practical implementation
2. Applied linear least squares to uncover the relationship between house size and price
3. Extended the approach to polynomial fitting for cyclical temperature data
4. Demonstrated non-linear least squares for exponential population growth
5. Conducted advanced analysis through residuals, confidence intervals, and goodness-of-fit metrics
6. Compared least squares with alternative methods

Key insights from our analysis:
- The relationship between house size and price follows a linear law: Price ≈ 0.10 × Size + 1.94 ($1000s)
- Monthly temperatures follow a 4th-degree polynomial pattern, reflecting seasonal variations
- Population growth follows an exponential law: P(t) = 1.69e^(0.013t), indicating a constant growth rate

The least squares method continues to be an indispensable tool in scientific research, engineering, economics, and many other fields. Its ability to extract meaningful patterns from noisy data makes it essential for understanding complex systems and making predictions based on observed behavior.
"""
doc.add_paragraph(conclusion_text)

# Add References
doc.add_heading('9. References', level=1)
references_text = """
1. Björck, Å. (1996). Numerical Methods for Least Squares Problems. Society for Industrial and Applied Mathematics.
2. Golub, G. H., & Van Loan, C. F. (2013). Matrix Computations (4th ed.). Johns Hopkins University Press.
3. Lawson, C. L., & Hanson, R. J. (1995). Solving Least Squares Problems. Society for Industrial and Applied Mathematics.
4. Nocedal, J., & Wright, S. J. (2006). Numerical Optimization (2nd ed.). Springer.
5. Press, W. H., Teukolsky, S. A., Vetterling, W. T., & Flannery, B. P. (2007). Numerical Recipes: The Art of Scientific Computing (3rd ed.). Cambridge University Press.
6. Wikipedia: Least Squares. https://en.wikipedia.org/wiki/Least_squares
7. NIST/SEMATECH e-Handbook of Statistical Methods. https://www.itl.nist.gov/div898/handbook/
"""
doc.add_paragraph(references_text)

# Add footer
section = doc.sections[0]
footer = section.footer
footer_text = footer.paragraphs[0]
footer_text.text = "Prepared by Prem Bahadur Katuwal (202424080129) for PhD-level Numerical Analysis"
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
doc.save('Least_Squares_Report_Prem_Katuwal.docx')
